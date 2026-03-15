"""
Internal QA audit as a Word document for human review.

One QA .docx per memo run (e.g. 1120.SR_preview_20260312_131547_QA.docx).
Sections: Header validation, Key Preview, Recent Execution, Appendix B/C/D,
Valuation basis, Investment View fact pack, Investment View grounding (sentence table).
Primary review format; JSON remains optional for debugging.
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.services.report_styling import (
    ACCENT, BODY, GRAY,
    SPACE_TINY, SPACE_SMALL, SPACE_MED,
    TITLE_PT, SECTION_PT, BODY_PT, SMALL_PT, SOURCE_PT,
    apply_section_margins, set_cell_shading, style_table_header_row, set_compact_row_height,
)


def _run(paragraph, text: str, bold: bool = False, size_pt: float = 8.5, color=None):
    r = paragraph.add_run(str(text) if text is not None else "")
    r.font.name = "Arial"
    r.font.size = Pt(size_pt)
    r.font.color.rgb = color or BODY
    r.font.bold = bold
    return r


def _cell_text(v) -> str:
    if v is None:
        return "—"
    if isinstance(v, (int, float)):
        return f"{v:,.2f}" if isinstance(v, float) and v != int(v) else str(int(v)) if isinstance(v, float) and v == int(v) else str(v)
    return str(v)[:200]


def build_qa_audit_docx(
    qa_audit: dict,
    memo_data: dict,
    inv_sentences: list[dict] | None = None,
    ticker: str = "",
    duplicate_screening_log: list[dict] | None = None,
) -> Document:
    """Build QA audit Word document. Does not save; caller saves to path."""
    doc = Document()
    for section in doc.sections:
        apply_section_margins(section)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(BODY_PT)
    style.font.color.rgb = BODY
    style.paragraph_format.space_after = SPACE_TINY
    style.paragraph_format.space_before = Pt(0)

    # Title
    p = doc.add_paragraph()
    _run(p, "QA Audit — Earnings Preview", bold=True, size_pt=TITLE_PT, color=ACCENT)
    _run(p, f"  {ticker}", size_pt=TITLE_PT, color=BODY)
    p = doc.add_paragraph()
    _run(p, f"Generated: {qa_audit.get('generated_at', '')}", size_pt=SMALL_PT, color=GRAY)
    # MS section suppression (better blank than wrong)
    if qa_audit.get("ms_section_suppressed_due_to_entity_mismatch") or qa_audit.get("ms_section_suppressed_due_to_contamination") or qa_audit.get("reused_default_payload_detected"):
        p = doc.add_paragraph()
        _run(p, "MarketScreener payload validation:", bold=True, size_pt=SMALL_PT)
        parts = []
        if qa_audit.get("ms_section_suppressed_due_to_missing_current_data"):
            parts.append("suppressed (missing current data)")
        if qa_audit.get("ms_section_suppressed_due_to_entity_mismatch"):
            parts.append("suppressed (entity mismatch)")
        if qa_audit.get("ms_section_suppressed_due_to_contamination"):
            parts.append("suppressed (cross-company contamination)")
        if qa_audit.get("reused_default_payload_detected"):
            parts.append("reused default payload detected")
        p = doc.add_paragraph()
        _run(p, "  ".join(parts), size_pt=SMALL_PT, color=GRAY)
        _run(p, f"  |  payload_entity_match={qa_audit.get('payload_entity_match', True)}  |  payload_source_ticker={qa_audit.get('payload_source_ticker', '')}", size_pt=SMALL_PT, color=GRAY)
    doc.add_paragraph()

    entries = qa_audit.get("entries") or []
    # Group entries by section for clearer layout
    by_section: dict[str, list[dict]] = {}
    for e in entries:
        sec = e.get("section", "other")
        by_section.setdefault(sec, []).append(e)

    # Single table columns
    col_headers = ["Section", "Field", "Displayed Value", "Source Value", "Source", "Formula / Basis", "Status", "Notes"]

    # 1. Header field validation
    _add_section_heading(doc, "Header field validation")
    header_entries = [e for e in entries if e.get("section") == "header"]
    _add_validation_table(doc, header_entries, col_headers, _run)

    # 2. Key Preview validation
    _add_section_heading(doc, "Key Preview validation")
    kp_entries = [e for e in entries if e.get("section") == "key_preview"]
    _add_validation_table(doc, kp_entries, col_headers, _run)

    # 3. Recent Execution validation
    _add_section_heading(doc, "Recent Execution validation")
    rec_entries = [e for e in entries if e.get("section") == "recent_execution"]
    _add_validation_table(doc, rec_entries, col_headers, _run)

    # 4. Appendix B surprise validation
    _add_section_heading(doc, "Appendix B surprise validation")
    appb_entries = [e for e in entries if e.get("section") == "appendix_b"]
    _add_validation_table(doc, appb_entries, col_headers, _run)

    # 5. Appendix C / D validation
    _add_section_heading(doc, "Appendix C / D validation")
    appc_entries = [e for e in entries if e.get("section") == "appendix_c"]
    appd_entries = [e for e in entries if e.get("section") == "appendix_d"]
    _add_validation_table(doc, appc_entries + appd_entries, col_headers, _run)

    # 6. Valuation basis (from appendix_d and street_snapshot)
    _add_section_heading(doc, "Valuation basis checks")
    street_entries = [e for e in entries if e.get("section") == "street_snapshot"]
    val_entries = [e for e in street_entries if e.get("field_name") == "pe"] + [e for e in appd_entries]
    _add_validation_table(doc, val_entries, col_headers, _run)

    # 7. Investment View fact pack
    _add_section_heading(doc, "Investment View fact pack")
    fact_pack = _get_fact_pack_summary(memo_data)
    for line in fact_pack:
        para = doc.add_paragraph(style="List Bullet")
        _run(para, line, size_pt=SOURCE_PT)
    doc.add_paragraph()

    # 8. Investment View grounding: sentence traceability (Sentence | Source type | Article headline | Status)
    _add_section_heading(doc, "Investment View sentence traceability")
    inv_trace = qa_audit.get("investment_view_sentences") or []
    if inv_trace:
        tbl = doc.add_table(rows=1 + len(inv_trace), cols=4)
        tbl.style = "Table Grid"
        style_table_header_row(tbl.rows[0], ["Sentence", "Source type", "Article headline", "Status"], _run)
        for i, row in enumerate(inv_trace):
            cells = tbl.rows[i + 1].cells
            cells[0].paragraphs[0].clear()
            _run(cells[0].paragraphs[0], row.get("sentence", "")[:400], size_pt=SOURCE_PT)
            cells[1].paragraphs[0].clear()
            _run(cells[1].paragraphs[0], row.get("source_type", "memo_fact"), size_pt=SOURCE_PT, color=ACCENT if row.get("source_type") == "recent_context_fact" else BODY)
            cells[2].paragraphs[0].clear()
            _run(cells[2].paragraphs[0], (row.get("article_headline") or "")[:80], size_pt=SOURCE_PT, color=GRAY)
            cells[3].paragraphs[0].clear()
            _run(cells[3].paragraphs[0], row.get("status", "kept"), size_pt=SOURCE_PT)
            set_compact_row_height(tbl.rows[i + 1], 16)
    elif inv_sentences:
        tbl = doc.add_table(rows=1 + len(inv_sentences), cols=3)
        tbl.style = "Table Grid"
        style_table_header_row(tbl.rows[0], ["Sentence", "Classification", "Status"], _run)
        for i, row in enumerate(inv_sentences):
            cells = tbl.rows[i + 1].cells
            cells[0].paragraphs[0].clear()
            _run(cells[0].paragraphs[0], row.get("sentence", "")[:400], size_pt=SOURCE_PT)
            cells[1].paragraphs[0].clear()
            _run(cells[1].paragraphs[0], row.get("classification", ""), size_pt=SOURCE_PT)
            cells[2].paragraphs[0].clear()
            _run(cells[2].paragraphs[0], row.get("status", ""), size_pt=SOURCE_PT, color=ACCENT if row.get("status") == "removed" else BODY)
            set_compact_row_height(tbl.rows[i + 1], 16)
    else:
        p = doc.add_paragraph()
        _run(p, "No Investment View sentences (e.g. fallback text used).", size_pt=SMALL_PT, color=GRAY)
    doc.add_paragraph()

    # Recent context checks (fail loudly if news found but not in DOCX or IV didn't cite)
    _add_section_heading(doc, "Recent context checks")
    for label, val in [
        ("recent_context_retrieved", qa_audit.get("recent_context_retrieved")),
        ("recent_context_has_valid_articles", qa_audit.get("recent_context_has_valid_articles")),
        ("recent_context_candidate_count", qa_audit.get("recent_context_candidate_count", 0)),
        ("recent_context_valid_count", qa_audit.get("recent_context_valid_count", 0)),
        ("recent_context_article_count", qa_audit.get("recent_context_article_count", 0)),
        ("recent_context_render_count", qa_audit.get("recent_context_render_count", 0)),
        ("recent_context_rendered", qa_audit.get("recent_context_rendered")),
        ("recent_context_failure_stage", qa_audit.get("recent_context_failure_stage", "")),
        ("investment_view_used_article_count", qa_audit.get("investment_view_used_article_count", 0)),
        ("investment_view_uses_recent_context", qa_audit.get("investment_view_uses_recent_context")),
        ("candidate_valid_basic", qa_audit.get("candidate_valid_basic")),
        ("candidate_has_date_before_enrichment", qa_audit.get("candidate_has_date_before_enrichment", 0)),
        ("candidate_has_extracted_fact", qa_audit.get("candidate_has_extracted_fact", 0)),
        ("final_article_valid_count", qa_audit.get("final_article_valid_count", 0)),
        ("date_parse_attempted", qa_audit.get("date_parse_attempted", 0)),
        ("date_parse_source", qa_audit.get("date_parse_source")),
        ("date_parse_success", qa_audit.get("date_parse_success", 0)),
        ("candidates_rejected_for_missing_date", qa_audit.get("candidates_rejected_for_missing_date", 0)),
        ("candidates_recovered_after_article_fetch", qa_audit.get("candidates_recovered_after_article_fetch", 0)),
    ]:
        p = doc.add_paragraph()
        _run(p, f"{label}: ", bold=True, size_pt=SOURCE_PT)
        if isinstance(val, bool):
            _run(p, "true" if val else "false", size_pt=SOURCE_PT, color=ACCENT if val else GRAY)
        else:
            _run(p, str(val), size_pt=SOURCE_PT, color=GRAY)
    headlines = qa_audit.get("recent_context_used_headlines") or qa_audit.get("investment_view_used_article_headlines") or []
    if headlines:
        p = doc.add_paragraph()
        _run(p, "Recent context used headlines: ", bold=True, size_pt=SOURCE_PT)
        for h in headlines[:5]:
            para = doc.add_paragraph(style="List Bullet")
            _run(para, (h or "")[:120], size_pt=SOURCE_PT)
    rejected = qa_audit.get("recent_context_rejected_reasons") or []
    if rejected:
        p = doc.add_paragraph()
        _run(p, "Rejected reasons (sample): ", bold=True, size_pt=SOURCE_PT)
        _run(p, ", ".join(rejected[:15]), size_pt=SOURCE_PT, color=GRAY)
    query_log = qa_audit.get("recent_context_query_log") or []
    if query_log:
        p = doc.add_paragraph()
        _run(p, "Query log: ", bold=True, size_pt=SOURCE_PT)
        for entry in query_log[:10]:
            para = doc.add_paragraph(style="List Bullet")
            _run(para, f"{entry.get('query', '')} [{entry.get('source', '')}] → {entry.get('count', 0)}", size_pt=SOURCE_PT, color=GRAY)
    enrichment_log = qa_audit.get("recent_context_enrichment_log") or []
    if enrichment_log:
        p = doc.add_paragraph()
        _run(p, "Enrichment log (date parse): ", bold=True, size_pt=SOURCE_PT)
        for entry in enrichment_log[:10]:
            para = doc.add_paragraph(style="List Bullet")
            _run(para, f"attempted={entry.get('date_parse_attempted')} source={entry.get('date_parse_source', '')} success={entry.get('date_parse_success')}", size_pt=SOURCE_PT, color=GRAY)
    articles_qa = qa_audit.get("recent_context_articles_qa") or []
    if articles_qa:
        p = doc.add_paragraph()
        _run(p, "Recent context articles (date_source, date_confidence, extracted_fact, relevance_reason): ", bold=True, size_pt=SOURCE_PT)
        for entry in articles_qa[:10]:
            headline = (entry.get("headline") or "")[:60]
            ds, dc = entry.get("date_source", ""), entry.get("date_confidence", "")
            fact = (entry.get("extracted_fact") or "")[:120]
            reason = entry.get("relevance_reason", "")
            para = doc.add_paragraph(style="List Bullet")
            _run(para, f"{headline} | source={ds} confidence={dc} | fact={fact} | reason={reason}", size_pt=SOURCE_PT, color=GRAY)
    rejected_top = qa_audit.get("rejected_candidates_top_10") or []
    if rejected_top:
        p = doc.add_paragraph()
        _run(p, "Top rejected candidates (debug): ", bold=True, size_pt=SOURCE_PT)
        for entry in rejected_top[:10]:
            para = doc.add_paragraph(style="List Bullet")
            _run(para, f"[{entry.get('provider', '')}] {(entry.get('headline') or '')[:50]} | search_card={entry.get('search_card_date_found')} fetch_ok={entry.get('article_fetch_succeeded')} selectors={entry.get('date_selectors_tried', '')} reason={entry.get('rejection_reason', '')}", size_pt=SOURCE_PT, color=GRAY)
    fail_reason = qa_audit.get("recent_context_render_failed_reason")
    if fail_reason:
        p = doc.add_paragraph()
        _run(p, "QA FAIL: ", bold=True, size_pt=SOURCE_PT, color=ACCENT)
        _run(p, fail_reason, size_pt=SOURCE_PT, color=ACCENT)
    elif qa_audit.get("recent_context_retrieved") and qa_audit.get("recent_context_rendered") is False:
        p = doc.add_paragraph()
        _run(p, "QA FAIL: recent news was found but not rendered in the memo DOCX.", size_pt=SOURCE_PT, color=ACCENT)
    doc.add_paragraph()

    # Duplicate screening log (news dedupe)
    if duplicate_screening_log:
        _add_section_heading(doc, "Duplicate screening log")
        tbl = doc.add_table(rows=1 + len(duplicate_screening_log), cols=4)
        tbl.style = "Table Grid"
        style_table_header_row(tbl.rows[0], ["Original headline", "Normalized headline", "Kept/Dropped", "Reason"], _run)
        for i, row in enumerate(duplicate_screening_log):
            cells = tbl.rows[i + 1].cells
            cells[0].paragraphs[0].clear()
            _run(cells[0].paragraphs[0], (row.get("original_headline") or "")[:100], size_pt=SOURCE_PT)
            cells[1].paragraphs[0].clear()
            _run(cells[1].paragraphs[0], (row.get("normalized_headline") or "")[:80], size_pt=SOURCE_PT)
            cells[2].paragraphs[0].clear()
            _run(cells[2].paragraphs[0], "kept" if row.get("kept") else "dropped", size_pt=SOURCE_PT)
            cells[3].paragraphs[0].clear()
            _run(cells[3].paragraphs[0], (row.get("reason") or "")[:60], size_pt=SOURCE_PT, color=GRAY)
            set_compact_row_height(tbl.rows[i + 1], 14)
        doc.add_paragraph()

    # Snapshot summary
    _add_section_heading(doc, "Snapshot summary")
    snap = qa_audit.get("snapshots_summary") or {}
    for name, info in snap.items():
        p = doc.add_paragraph()
        _run(p, f"{name}: ", bold=True, size_pt=SOURCE_PT)
        _run(p, str(info), size_pt=SOURCE_PT, color=GRAY)
    return doc


def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SPACE_SMALL
    _run(p, title, bold=True, size_pt=SECTION_PT, color=ACCENT)
    p.paragraph_format.space_after = SPACE_TINY


def _add_validation_table(doc: Document, entries: list[dict], col_headers: list[str], run_fn) -> None:
    if not entries:
        p = doc.add_paragraph()
        run_fn(p, "No entries.", size_pt=SMALL_PT, color=GRAY)
        return
    tbl = doc.add_table(rows=1 + len(entries), cols=len(col_headers))
    tbl.style = "Table Grid"
    style_table_header_row(tbl.rows[0], col_headers, run_fn)
    for i, e in enumerate(entries):
        row_cells = tbl.rows[i + 1].cells
        row_cells[0].paragraphs[0].clear()
        run_fn(row_cells[0].paragraphs[0], e.get("section", ""), size_pt=SOURCE_PT)
        row_cells[1].paragraphs[0].clear()
        run_fn(row_cells[1].paragraphs[0], e.get("field_name", ""), size_pt=SOURCE_PT)
        row_cells[2].paragraphs[0].clear()
        run_fn(row_cells[2].paragraphs[0], _cell_text(e.get("displayed_value")), size_pt=SOURCE_PT)
        row_cells[3].paragraphs[0].clear()
        run_fn(row_cells[3].paragraphs[0], _cell_text(e.get("raw_source_value")), size_pt=SOURCE_PT)
        row_cells[4].paragraphs[0].clear()
        run_fn(row_cells[4].paragraphs[0], (e.get("source_name") or "") + " " + (e.get("source_url") or ""), size_pt=SOURCE_PT)
        row_cells[5].paragraphs[0].clear()
        run_fn(row_cells[5].paragraphs[0], e.get("formula_used") or "—", size_pt=SOURCE_PT)
        row_cells[6].paragraphs[0].clear()
        run_fn(row_cells[6].paragraphs[0], e.get("status", ""), size_pt=SOURCE_PT)
        row_cells[7].paragraphs[0].clear()
        run_fn(row_cells[7].paragraphs[0], (e.get("notes") or "")[:150], size_pt=SOURCE_PT, color=GRAY)
        set_compact_row_height(tbl.rows[i + 1], 14)
    doc.add_paragraph()


def _get_fact_pack_summary(memo_data: dict) -> list[str]:
    """Bullet summary of fact pack for QA docx."""
    lines = []
    entity = memo_data.get("entity") or {}
    header = memo_data.get("header") or {}
    key_preview = memo_data.get("key_preview") or []
    rec = memo_data.get("recent_execution") or {}

    def _v(f):
        if not isinstance(f, dict):
            return None
        return f.get("display_value") if f.get("display_value") is not None else f.get("value")

    lines.append(f"Company: {entity.get('company_name')} ({entity.get('ticker')})")
    lines.append(f"Expected report date: {_v(header.get('expected_report_date'))}")
    lines.append(f"Recommendation: {_v(header.get('recommendation'))}; Analysts: {_v(header.get('analyst_count'))}")
    lines.append(f"Price (consensus): {_v(header.get('consensus_page_price'))}; Target: {_v(header.get('average_target_price'))}; Upside: {_v(header.get('upside_pct'))}%")
    for row in key_preview[:5]:
        k = row.get("metric_key", "")
        c, pr, sy = _v(row.get("consensus")), _v(row.get("prior_actual")), _v(row.get("same_q_prior_yr"))
        lines.append(f"Key preview {k}: consensus={c}, prior={pr}, same_ly={sy}")
    lines.append(f"Avg revenue surprise: {_v(rec.get('avg_revenue_surprise_pct'))}%; EPS: {_v(rec.get('avg_eps_surprise_pct'))}%")
    lines.append(f"Revenue surprise history: {len(rec.get('revenue_surprise_history') or [])} quarters")
    return lines


def write_qa_audit_docx(
    qa_audit: dict,
    memo_data: dict,
    path: Path,
    inv_sentences: list[dict] | None = None,
    ticker: str = "",
    duplicate_screening_log: list[dict] | None = None,
) -> None:
    """Build and save QA audit document to path."""
    doc = build_qa_audit_docx(
        qa_audit, memo_data,
        inv_sentences=inv_sentences, ticker=ticker,
        duplicate_screening_log=duplicate_screening_log,
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
